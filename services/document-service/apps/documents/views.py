from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.http import HttpResponse
from django.conf import settings
from io import BytesIO
from minio import Minio
from minio.error import S3Error
import jwt
import json
import base64
from decouple import config
from .models import Document, DocumentSignature
from .serializers import DocumentSerializer, DocumentSignatureSerializer
import hashlib
import os
from urllib.request import Request, urlopen
from urllib.error import HTTPError, URLError
from datetime import datetime

WORD_MIME_TYPES = {
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/msword',
}


def _build_signed_docx(doc_bytes: bytes, signer_name: str, sig_type: str,
                        sig_data: str, stamp_type: str, date_str: str) -> bytes:
    """Append a professional LAWYER'S SIGNATURE section to a Word document."""
    from docx import Document as WordDoc
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    STAMP_COLORS = {
        'reviewed':     RGBColor(0x1e, 0x9e, 0x40),
        'approved':     RGBColor(0x14, 0x70, 0xc4),
        'certified':    RGBColor(0x70, 0x10, 0xb4),
        'confidential': RGBColor(0xc4, 0x14, 0x14),
        'court_filed':  RGBColor(0xc4, 0x74, 0x10),
        'original_copy':RGBColor(0x10, 0x5c, 0x9e),
    }

    wdoc = WordDoc(BytesIO(doc_bytes))

    # separator
    wdoc.add_paragraph('').add_run('─' * 60).font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

    # heading
    h = wdoc.add_paragraph()
    hr = h.add_run("LAWYER'S SIGNATURE")
    hr.bold = True
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # signature content
    if sig_type == 'stamp' and stamp_type:
        color = STAMP_COLORS.get(stamp_type, RGBColor(0x44, 0x44, 0x44))
        label = stamp_type.replace('_', ' ').upper()
        tbl = wdoc.add_table(rows=1, cols=1)
        tbl.style = 'Table Grid'
        cell = tbl.rows[0].cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = color

    elif sig_type == 'typed':
        p = wdoc.add_paragraph()
        run = p.add_run(sig_data)
        run.italic = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x5e)

    elif sig_type == 'draw':
        try:
            raw = sig_data.split(',', 1)[-1]
            img_bytes = base64.b64decode(raw)
            p = wdoc.add_paragraph()
            p.add_run().add_picture(BytesIO(img_bytes), width=Inches(2.2))
        except Exception as exc:
            p = wdoc.add_paragraph()
            p.add_run(f'[Signature image — {exc}]')

    # footer line
    footer = wdoc.add_paragraph()
    footer.add_run(f'Signed by: {signer_name or "Lawyer"}   |   Date: {date_str}').font.size = Pt(8)

    out = BytesIO()
    wdoc.save(out)
    return out.getvalue()


def extract_user_id_from_token(request):
    """Extract user_id UUID from JWT token"""
    auth_header = request.META.get('HTTP_AUTHORIZATION', '')
    if auth_header.startswith('Bearer '):
        try:
            token = auth_header.split(' ')[1]
            signing_key = settings.SIMPLE_JWT.get('SIGNING_KEY', settings.SECRET_KEY)
            payload = jwt.decode(token, signing_key, algorithms=['HS256'], options={'verify_aud': False})
            return payload.get('user_id')
        except Exception:
            pass
    return str(request.user.id)


STAFF_ROLES = {'lawyer', 'firm_admin', 'firm-admin', 'partner', 'associate', 'secretary', 'managing_partner'}


def extract_auth_payload(request):
    auth_header = request.META.get('HTTP_AUTHORIZATION', '')
    if not auth_header.startswith('Bearer '):
        return {}
    token = auth_header.split(' ')[1]
    try:
        signing_key = settings.SIMPLE_JWT.get('SIGNING_KEY', settings.SECRET_KEY)
        return jwt.decode(token, signing_key, algorithms=['HS256'], options={'verify_aud': False})
    except Exception:
        return {}


def get_minio_client():
    endpoint = settings.MINIO_ENDPOINT
    return Minio(
        endpoint,
        access_key=settings.MINIO_ACCESS_KEY,
        secret_key=settings.MINIO_SECRET_KEY,
        secure=False,
    )


def ensure_bucket_exists(client):
    bucket_name = settings.MINIO_BUCKET_NAME
    if not client.bucket_exists(bucket_name):
        client.make_bucket(bucket_name)


def is_internal_request(request):
    return request.META.get('HTTP_X_INTERNAL_API_KEY') == config('INTERNAL_API_KEY', default='dev-internal-key')


# Returns 'ok', 'forbidden', 'unauthorized', or 'service_error'
def can_access_case(request, case_id):
    if is_internal_request(request):
        return 'ok'

    auth_header = request.META.get('HTTP_AUTHORIZATION', '')
    if not auth_header.startswith('Bearer '):
        return 'unauthorized'

    base_url = settings.CASE_SERVICE_URL.rstrip('/')
    url = f'{base_url}/{case_id}/'
    try:
        req = Request(url, headers={'Authorization': auth_header})
        with urlopen(req, timeout=5) as response:
            return 'ok' if response.status == 200 else 'forbidden'
    except HTTPError as e:
        return 'forbidden' if e.code in (403, 404) else 'service_error'
    except (URLError, ValueError):
        return 'service_error'


def _access_error_response(access_result):
    """Convert a non-ok access result into a DRF Response."""
    if access_result == 'service_error':
        return Response(
            {'error': 'Unable to verify case access. The case service is temporarily unavailable. Please try again shortly.'},
            status=status.HTTP_503_SERVICE_UNAVAILABLE,
        )
    return Response(
        {'error': 'You do not have access to this matter. Please check that the matter exists and is assigned to you.'},
        status=status.HTTP_403_FORBIDDEN,
    )


def _get_case_info(case_id, auth_header):
    """Fetch case details from case service. Returns dict or None."""
    base_url = settings.CASE_SERVICE_URL.rstrip('/')
    url = f'{base_url}/{case_id}/'
    try:
        req = Request(url, headers={'Authorization': auth_header})
        with urlopen(req, timeout=5) as response:
            return json.loads(response.read().decode('utf-8'))
    except Exception:
        return None


def _notify_document_uploaded(case_id, document, uploader_role, auth_header):
    """Fire-and-forget: notify the relevant party when a document is uploaded."""
    try:
        case = _get_case_info(case_id, auth_header)
        if not case:
            return

        if uploader_role in STAFF_ROLES:
            recipient_id = case.get('client_id')
            title_en = 'New Document Added by Your Lawyer'
            title_fr = 'Nouveau Document Ajouté par Votre Avocat'
            msg_en = f'A new document "{document.filename}" has been added to your case "{case.get("title", "")}".'
            msg_fr = f'Un nouveau document "{document.filename}" a été ajouté à votre dossier "{case.get("title", "")}".'
        else:
            recipient_id = case.get('assigned_lawyer_id')
            title_en = 'New Document Uploaded by Client'
            title_fr = 'Nouveau Document Téléchargé par le Client'
            msg_en = f'Your client uploaded "{document.filename}" to case "{case.get("title", "")}".'
            msg_fr = f'Votre client a téléchargé "{document.filename}" dans le dossier "{case.get("title", "")}".'

        if not recipient_id:
            return

        notify_url = getattr(settings, 'NOTIFICATION_SERVICE_URL', 'http://notification-service:8006').rstrip('/') + '/api/v1/notifications/internal/create/'
        internal_key = config('INTERNAL_API_KEY', default='dev-internal-key')

        payload = json.dumps({
            'user_id': str(recipient_id),
            'event_type': 'document_uploaded',
            'title_en': title_en,
            'title_fr': title_fr,
            'message_en': msg_en,
            'message_fr': msg_fr,
            'metadata': {
                'document_id': str(document.id),
                'case_id': str(case_id),
                'filename': document.filename,
            },
        }).encode('utf-8')

        req = Request(
            notify_url,
            data=payload,
            headers={'X-Internal-Key': internal_key, 'Content-Type': 'application/json'},
            method='POST',
        )
        with urlopen(req, timeout=5):
            pass
    except Exception as exc:
        print(f'[notify_document_uploaded] non-fatal: {exc}')


class DocumentUploadView(APIView):
    """Upload a document for a case"""

    def post(self, request, case_id):
        access = can_access_case(request, case_id)
        if access != 'ok':
            return _access_error_response(access)

        file_obj = request.FILES.get('file')
        doc_type = request.data.get('type', 'other')
        user_id = extract_user_id_from_token(request)

        if not file_obj:
            return Response({'error': 'file required'}, status=status.HTTP_400_BAD_REQUEST)

        parent_id = request.data.get('parent_document_id') or None
        version = 1
        if parent_id:
            try:
                parent = Document.objects.get(id=parent_id)
                version = parent.version + 1
            except Document.DoesNotExist:
                parent_id = None

        document = Document.objects.create(
            case_id=case_id,
            uploader_id=user_id,
            filename=file_obj.name,
            document_type=doc_type,
            file_size=file_obj.size,
            mime_type=file_obj.content_type,
            status='pending_scan',
            version=version,
            parent_document_id=parent_id,
        )

        password = request.data.get('password')
        if password:
            salt = os.urandom(16).hex()
            digest = hashlib.sha256((salt + password).encode('utf-8')).hexdigest()
            document.password_salt = salt
            document.password_hash = digest
            document.save(update_fields=['password_salt', 'password_hash'])

        try:
            client = get_minio_client()
            ensure_bucket_exists(client)
            object_name = f'cases/{case_id}/{document.id}'
            file_bytes = file_obj.read()
            file_obj.seek(0)
            client.put_object(
                settings.MINIO_BUCKET_NAME,
                object_name,
                BytesIO(file_bytes),
                len(file_bytes),
                content_type=file_obj.content_type or 'application/octet-stream',
            )
            document.minio_path = object_name
            document.status = 'stored'
            document.save(update_fields=['minio_path', 'status', 'updated_at'])
        except Exception as exc:
            document.delete()
            return Response({'error': f'failed to store document: {exc}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        auth_header = request.META.get('HTTP_AUTHORIZATION', '')
        caller_role = extract_auth_payload(request).get('role', '')
        _notify_document_uploaded(case_id, document, caller_role, auth_header)

        return Response(
            DocumentSerializer(document).data,
            status=status.HTTP_201_CREATED
        )


class DocumentDownloadView(APIView):
    """Download a document from MinIO."""

    authentication_classes = []
    permission_classes = []

    def get(self, request, document_id):
        is_internal = is_internal_request(request)
        caller_payload = extract_auth_payload(request) if not is_internal else {}
        caller_id = caller_payload.get('user_id')
        caller_role = caller_payload.get('role', '')

        if not is_internal and not caller_id:
            return Response({'error': 'forbidden'}, status=status.HTTP_403_FORBIDDEN)

        try:
            document = Document.objects.get(id=document_id)
        except Document.DoesNotExist:
            return Response({'error': 'not found'}, status=status.HTTP_404_NOT_FOUND)

        # Case service is the single source of truth for access — it handles
        # clients, assigned lawyers, other firm lawyers, secretaries, and firm admins.
        access = can_access_case(request, document.case_id)
        if access != 'ok':
            return _access_error_response(access)

        if not document.minio_path:
            return Response({'error': 'document not stored'}, status=status.HTTP_404_NOT_FOUND)

        # Firm staff bypass the password requirement — they manage the case
        # and should not be blocked by client-set document passwords.
        is_staff = caller_role in STAFF_ROLES
        if document.password_hash and not is_internal and not is_staff:
            provided = request.META.get('HTTP_X_DOCUMENT_PASSWORD') or request.GET.get('password')
            if not provided:
                return Response({'error': 'password required'}, status=status.HTTP_403_FORBIDDEN)
            check = hashlib.sha256((document.password_salt + provided).encode('utf-8')).hexdigest()
            if check != document.password_hash:
                return Response({'error': 'invalid password'}, status=status.HTTP_403_FORBIDDEN)

        try:
            client = get_minio_client()
            response = client.get_object(settings.MINIO_BUCKET_NAME, document.minio_path)
            file_bytes = response.read()
            response.close()
            response.release_conn()
        except S3Error as exc:
            return Response({'error': f'minio error: {exc}'}, status=status.HTTP_404_NOT_FOUND)

        http_response = HttpResponse(file_bytes, content_type=document.mime_type or 'application/octet-stream')
        http_response['Content-Disposition'] = f'attachment; filename="{document.filename}"'
        return http_response


class DocumentListView(APIView):
    """List documents for a case"""

    def get(self, request, case_id):
        access = can_access_case(request, case_id)
        if access != 'ok':
            return _access_error_response(access)

        documents = Document.objects.filter(case_id=case_id).prefetch_related('signatures')
        serializer = DocumentSerializer(documents, many=True)
        return Response({
            'count': documents.count(),
            'results': serializer.data
        })


class SignDocumentView(APIView):
    """POST /api/v1/documents/<document_id>/sign/ — lawyer/staff only."""

    def post(self, request, document_id):
        try:
            document = Document.objects.get(id=document_id)
        except Document.DoesNotExist:
            return Response({'error': 'not found'}, status=status.HTTP_404_NOT_FOUND)

        access = can_access_case(request, document.case_id)
        if access != 'ok':
            return _access_error_response(access)

        payload = extract_auth_payload(request)
        caller_role = payload.get('role', '')
        if caller_role not in STAFF_ROLES:
            return Response({'error': 'Only lawyers and firm staff may sign documents.'}, status=status.HTTP_403_FORBIDDEN)

        sig_type = request.data.get('signature_type', '')
        sig_data = request.data.get('signature_data', '')
        stamp    = request.data.get('stamp_type', '')
        name     = request.data.get('signer_name', '')
        signer   = str(payload.get('user_id', ''))

        if sig_type not in ('draw', 'typed', 'stamp'):
            return Response({'error': 'signature_type must be draw, typed, or stamp'}, status=status.HTTP_400_BAD_REQUEST)
        if not sig_data:
            return Response({'error': 'signature_data required'}, status=status.HTTP_400_BAD_REQUEST)

        sig = DocumentSignature.objects.create(
            document=document,
            signer_id=signer,
            signer_name=name,
            signature_type=sig_type,
            signature_data=sig_data,
            stamp_type=stamp,
        )

        # For Word documents, build a signed .docx on the backend and
        # store it as a new version (PDF/image baking is handled client-side).
        signed_doc_id = None
        if document.mime_type in WORD_MIME_TYPES and document.minio_path:
            try:
                minio_client = get_minio_client()
                resp = minio_client.get_object(settings.MINIO_BUCKET_NAME, document.minio_path)
                doc_bytes = resp.read()
                resp.close(); resp.release_conn()

                date_str = datetime.utcnow().strftime('%d %b %Y')
                signed_bytes = _build_signed_docx(doc_bytes, name, sig_type, sig_data, stamp, date_str)

                base_name = document.filename
                if base_name.lower().endswith('.docx'):
                    new_name = base_name[:-5] + f'_signed_v{document.version + 1}.docx'
                elif base_name.lower().endswith('.doc'):
                    new_name = base_name[:-4] + f'_signed_v{document.version + 1}.doc'
                else:
                    new_name = base_name + f'_signed_v{document.version + 1}'

                new_doc = Document.objects.create(
                    case_id=document.case_id,
                    uploader_id=signer,
                    filename=new_name,
                    document_type=document.document_type,
                    file_size=len(signed_bytes),
                    mime_type=document.mime_type,
                    status='pending_scan',
                    version=document.version + 1,
                    parent_document_id=str(document.id),
                )
                obj_name = f'cases/{document.case_id}/{new_doc.id}'
                minio_client.put_object(
                    settings.MINIO_BUCKET_NAME, obj_name,
                    BytesIO(signed_bytes), len(signed_bytes),
                    content_type=document.mime_type,
                )
                new_doc.minio_path = obj_name
                new_doc.status = 'stored'
                new_doc.save(update_fields=['minio_path', 'status', 'updated_at'])
                signed_doc_id = str(new_doc.id)
            except Exception as exc:
                print(f'[sign_word_doc] non-fatal: {exc}')

        data = DocumentSignatureSerializer(sig).data
        if signed_doc_id:
            data['signed_version_id'] = signed_doc_id
        return Response(data, status=status.HTTP_201_CREATED)
